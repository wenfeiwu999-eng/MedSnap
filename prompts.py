# -*- coding: utf-8 -*-
"""MedSnap AI prompts"""

# ========== 内置模板Prompt定义 ==========

PROMPT_DOCTOR_MEDICAL_RECORD = """你是临床医生数据提取专家。请仔细识别该病历图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "床号": "", "住院号": ""
  },
  "chief_complaint": "",
  "present_illness": "",
  "past_history": "",
  "personal_history": "",
  "family_history": "",
  "physical_exam": {
    "体温": null, "脉搏": null, "呼吸": null, "血压": "",
    "一般情况": "", "专科检查": ""
  },
  "diagnosis": [
    {"诊断名称": "", "ICD10编码": ""}
  ],
  "treatment_plan": {
    "药物治疗": "", "手术治疗": "", "其他治疗": ""
  },
  "surgery_record": "",
  "discharge_summary": "",
  "confidence": {}
}

## 提取规则
1. 年龄提取纯数字（如"56岁"→56）。
2. 诊断需尽量识别ICD-10编码（如 E11.9 2型糖尿病、I10 高血压病）。
3. 治疗方案区分药物治疗、手术治疗、其他治疗。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。未识别到的字段置信度为0。
5. 手写体请尽力识别，无法确认的字用?标记。

只输出JSON，不要输出任何其他内容。"""

PROMPT_DOCTOR_LAB_RESULTS = """你是临床检验数据提取专家。请仔细识别该检查检验结果图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "住院号": ""
  },
  "report_info": {
    "报告类型": "", "检查日期": "", "报告日期": ""
  },
  "lab_tests": [
    {
      "项目名称": "", "英文缩写": "", "数值": null,
      "单位": "", "参考范围": "", "异常标注": ""
    }
  ],
  "confidence": {}
}

## 提取规则
1. 识别常见医学缩写：WBC、RBC、PLT、Hb、ALT、AST、Cr、BUN、GLU、TC、TG、HDL-C、LDL-C、UA、CRP、ESR、HbA1c、TSH、FT3、FT4、Na、K、Ca、Cl、PT、APTT、INR、D-Dimer、AFP、CEA等。
2. 异常标注用"↑"(偏高)/"↓"(偏低)/"正常"。
3. 数值标准化为数字格式。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

## 医疗专业词汇参考
- 血常规：WBC、RBC、PLT、Hb、HCT、MCV、MCH、MCHC、RDW
- 肝功能：ALT、AST、GGT、ALP、TBIL、DBIL、TP、ALB
- 肾功能：Cr、BUN、UA、Cys-C、eGFR
- 血脂：TC、TG、HDL-C、LDL-C、ApoA1、ApoB
- 血糖：GLU、FPG、2hPG、HbA1c、OGTT
- 凝血：PT、APTT、TT、FIB、INR、D-Dimer
- 电解质：Na、K、Ca、Cl、Mg、P
- 炎症指标：CRP、PCT、ESR、IL-6
- 肿瘤标志物：AFP、CEA、CA125、CA199、CA153、PSA

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_ADMISSION = """你是护理评估数据提取专家。请仔细识别该入院护理评估表图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "床号": "", "住院号": "", "入院日期": ""
  },
  "vital_signs": {
    "体温": null, "脉搏": null, "呼吸": null, "血压_收缩压": null, "血压_舒张压": null,
    "血氧饱和度": null, "身高_cm": null, "体重_kg": null
  },
  "assessment": {
    "意识状态": "", "精神状态": "", "皮肤完整性": "", "皮肤异常描述": "",
    "营养状况": "", "饮食类型": "", "排便情况": "", "排尿情况": "",
    "睡眠情况": "", "活动能力": "", "自理能力初筛": "",
    "跌倒风险初筛": "", "压疮风险初筛": "", "疼痛评分": null,
    "过敏史": "", "特殊用药": ""
  },
  "nursing_diagnosis": [],
  "nursing_plan": "",
  "confidence": {}
}

## 提取规则
1. 生命体征数值标准化为纯数字。血压格式拆分为收缩压和舒张压。
2. 意识状态：清醒/嗜睡/昏睡/浅昏迷/深昏迷。
3. 自理能力/跌倒/压疮风险初筛：识别勾选框或评分。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_BARTHEL = """你是护理评估数据提取专家。请仔细识别该Barthel自理能力指数量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "barthel_items": [
    {"项目": "进食", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "洗澡", "评分": null, "评分标准": "5=自理 0=需帮助", "备注": ""},
    {"项目": "修饰", "评分": null, "评分标准": "5=自理 0=需帮助", "备注": ""},
    {"项目": "穿衣", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "控制大便", "评分": null, "评分标准": "10=可控 5=偶有失禁 0=失禁", "备注": ""},
    {"项目": "控制小便", "评分": null, "评分标准": "10=可控 5=偶有失禁 0=失禁", "备注": ""},
    {"项目": "如厕", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "床椅转移", "评分": null, "评分标准": "15=自理 10=少量帮助 5=较大帮助 0=完全依赖", "备注": ""},
    {"项目": "平地行走", "评分": null, "评分标准": "15=自行45m 10=在帮助下45m 5=轮椅45m 0=不能", "备注": ""},
    {"项目": "上下楼梯", "评分": null, "评分标准": "10=自理 5=需帮助 0=不能", "备注": ""}
  ],
  "total_score": null,
  "dependency_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字。
2. 总分范围0-100。依赖等级判定：100=自理、61-99=轻度依赖、41-60=中度依赖、≤40=重度依赖。
3. 如图中有勾选标记，按勾选对应的分值提取。
4. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_MORSE = """你是护理评估数据提取专家。请仔细识别该Morse跌倒风险评估量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "morse_items": [
    {"项目": "跌倒史", "评分": null, "评分标准": "是=25 否=0", "选项": ""},
    {"项目": "继发诊断", "评分": null, "评分标准": "是=15 否=0", "选项": ""},
    {"项目": "步行辅助", "评分": null, "评分标准": "卧床/护士协助=0 拐杖/助行器/轮椅=15 扶家具行走=30", "选项": ""},
    {"项目": "静脉输液/肝素锁", "评分": null, "评分标准": "是=20 否=0", "选项": ""},
    {"项目": "步态", "评分": null, "评分标准": "正常/卧床/不能活动=0 虚弱=10 损伤=20", "选项": ""},
    {"项目": "认知状态", "评分": null, "评分标准": "能正确认识自身活动能力=0 高估/忘记限制=15", "选项": ""}
  ],
  "total_score": null,
  "risk_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字。
2. 风险判定：0-24=低风险、25-44=中风险、≥45=高风险。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_BRADEN = """你是护理评估数据提取专家。请仔细识别该Braden压疮风险评估量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "braden_items": [
    {"项目": "感知能力", "评分": null, "评分标准": "1=完全受限 2=非常受限 3=轻度受限 4=未受损", "备注": ""},
    {"项目": "潮湿程度", "评分": null, "评分标准": "1=持续潮湿 2=非常潮湿 3=偶尔潮湿 4=很少潮湿", "备注": ""},
    {"项目": "活动能力", "评分": null, "评分标准": "1=卧床 2=坐椅 3=偶尔步行 4=经常步行", "备注": ""},
    {"项目": "移动能力", "评分": null, "评分标准": "1=完全不能 2=严重受限 3=轻度受限 4=不受限", "备注": ""},
    {"项目": "营养摄取", "评分": null, "评分标准": "1=非常差 2=可能不足 3=足够 4=良好", "备注": ""},
    {"项目": "摩擦力和剪切力", "评分": null, "评分标准": "1=存在问题 2=潜在问题 3=不存在问题", "备注": ""}
  ],
  "total_score": null,
  "risk_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字（1-4分，摩擦力1-3分）。
2. 总分6-23分。风险判定：≤9=极高风险、10-12=高风险、13-14=中度风险、15-18=低风险、≥19=无风险。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_PAIN = """你是护理评估数据提取专家。请仔细识别该疼痛评估记录图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "pain_assessment": {
    "疼痛部位": "",
    "疼痛性质": "",
    "NRS评分": null,
    "VAS评分": null,
    "疼痛频率": "",
    "持续时间": "",
    "加重因素": "",
    "缓解因素": "",
    "对睡眠影响": "",
    "对日常活动影响": "",
    "当前镇痛措施": "",
    "镇痛效果": ""
  },
  "confidence": {}
}

## 提取规则
1. NRS评分0-10（0=无痛，10=最剧烈疼痛）。VAS评分0-10。
2. 疼痛性质：锐痛/钝痛/刺痛/胀痛/灼痛/绞痛等。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_RECORD = """你是护理记录数据提取专家。请仔细识别该护理记录单图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "记录日期": ""
  },
  "vital_signs_records": [
    {"时间": "", "体温": null, "脉搏": null, "呼吸": null, "血压": "", "血氧": null, "备注": ""}
  ],
  "medication_execution": [
    {"时间": "", "医嘱内容": "", "执行情况": "", "执行人": ""}
  ],
  "nursing_measures": [
    {"时间": "", "护理措施": "", "患者反应": "", "记录人": ""}
  ],
  "handover_notes": "",
  "confidence": {}
}

## 提取规则
1. 时间格式统一为HH:MM。
2. 生命体征数值标准化为数字。
3. confidence字段：对每个记录项给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_RESEARCHER = """你是临床科研数据提取专家。请仔细识别该病历图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "demographics": {
    "姓名": "", "性别": "", "年龄": null, "身高_cm": null, "体重_kg": null,
    "BMI": null, "婚姻状况": "", "职业": "", "民族": "",
    "吸烟史": "", "饮酒史": "", "过敏史": ""
  },
  "lab_tests": [
    {"项目名称": "", "英文缩写": "", "数值": null, "单位": "", "参考范围": "", "异常标注": ""}
  ],
  "treatment": {
    "入院时间": "", "出院时间": "", "住院天数": null, "科室": "",
    "主诊断": "", "主诊断ICD10编码": "", "其他诊断": [],
    "手术操作": "", "治疗方案": "", "出院医嘱": ""
  },
  "confidence": {
    "demographics": {}, "lab_tests": {}, "treatment": {}
  }
}

## 提取规则
1. 年龄提取纯数字（如"56岁"→56）；BMI如未给出尝试计算。
2. 实验室检查识别常见缩写：WBC、RBC、PLT、Hb、ALT、AST、Cr、BUN、GLU、TC、TG、HDL-C、LDL-C、UA、CRP、ESR、HbA1c、TSH、FT3、FT4、Na、K、Ca、Cl。异常标注"↑"/"↓"/"正常"。
3. 诊疗资料尽量识别ICD-10编码；治疗方案区分药物/手术/其他。
4. 置信度0-1，未识别到的字段为0。

## 医疗专业词汇参考
- 血常规：WBC、RBC、PLT、Hb、HCT、MCV、MCH、MCHC、RDW
- 肝功能：ALT、AST、GGT、ALP、TBIL、DBIL、TP、ALB
- 肾功能：Cr、BUN、UA、Cys-C、eGFR
- 血脂：TC、TG、HDL-C、LDL-C、ApoA1、ApoB
- 血糖：GLU、FPG、2hPG、HbA1c、OGTT
- 凝血：PT、APTT、TT、FIB、INR、D-Dimer
- 甲功：TSH、FT3、FT4、T3、T4
- 电解质：Na、K、Ca、Cl、Mg、P
- 炎症指标：CRP、PCT、ESR、IL-6
- 肿瘤标志物：AFP、CEA、CA125、CA199、CA153、PSA

只输出JSON，不要输出任何其他内容。"""

# 护士自定义模板的Prompt生成框架
NURSE_CUSTOM_PROMPT_TEMPLATE = """你是护理数据结构化提取专家。请仔细识别该医疗文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 量表/问卷识别规则（重要）
如果文档包含评分量表或问卷（如Likert量表、Barthel指数、Morse跌倒量表等）：
- 每个题目的勾选框（checkbox/圈选/打勾）代表该题目的得分
- 观察被勾选/圈选的选项在量表中的位置编号，将其转换为对应数值
- 例如：5点量表中勾选第4个框 → 值为"4"，勾选第1个框 → 值为"1"
- 字段值必须是数字（如"4"），不要写描述文字

## 返回格式要求
请直接输出以下合法JSON：

{{
  "patient_info": {{
    "姓名": "", "床号": "", "科室": "", "日期": ""
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与这些字段最相关的具体信息：{field_names}。
2. 对于量表/问卷评分项，提取被勾选选项对应的数值（不是文字描述）。
{score_rule}
3. confidence字段：给每个已提取字段赋予0-1之间的置信度。

只返回JSON，不要包含任何其他内容。"""


# 医生自定义模板的Prompt生成框架
DOCTOR_CUSTOM_PROMPT_TEMPLATE = """你是临床数据结构化提取专家。请仔细识别该医疗文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 量表/问卷识别规则（重要）
如果文档包含评分量表或问卷（如ICCAS、PHQ-9、GAD-7等Likert量表）：
- 每个题目的勾选框（checkbox/圈选/打勾）代表该题目的得分
- 观察被勾选/圈选的选项在量表中的位置编号，将其转换为对应数值
- 例如：5点量表中勾选第4个框 → 值为"4"，勾选第1个框 → 值为"1"
- 字段值必须是数字（如"4"），不要写描述文字

## 返回格式要求
请直接输出以下合法JSON：

{{
  "patient_info": {{
    "姓名": "", "性别": "", "年龄": null, "科室": "", "诊断": "", "住院号": ""
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与这些字段最相关的具体信息：{field_names}。
2. 对于量表/问卷评分项，提取被勾选选项对应的数值（不是文字描述）。
3. 诊断尽量提取完整名称，并尽可能识别ICD-10编码。
4. 手写内容尽量识别，无法确认的用"?"标注。
5. confidence字段：给每个已提取字段赋予0-1之间的置信度。

只返回JSON，不要包含任何其他内容。"""


# 科研自定义模板的Prompt生成框架
RESEARCHER_CUSTOM_PROMPT_TEMPLATE = """你是临床数据结构化提取专家。请仔细识别该医疗文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 量表/问卷识别规则（重要）
如果文档包含评分量表或问卷（如ICCAS、PHQ-9、GAD-7等Likert量表）：
- 每个题目的勾选框（checkbox/圈选/打勾）代表该题目的得分
- 观察被勾选/圈选的选项在量表中的位置编号，将其转换为对应数值
- 例如：5点量表中勾选第4个框 → 值为"4"，勾选第1个框 → 值为"1"
- 字段值必须是数字（如"4"），不要写描述文字

## 返回格式要求
请直接输出以下合法JSON：

{{
  "demographics": {{
    "姓名": "", "性别": "", "年龄": null
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与这些字段最相关的具体信息：{field_names}。
2. 对于量表/问卷评分项，提取被勾选选项对应的数值（不是文字描述）。
3. 数值型数据提取为数字，文本型数据提取为字符串。
4. 日期格式统一为YYYY-MM-DD。
5. confidence字段：给每个已提取字段赋予0-1之间的置信度。

只返回JSON，不要包含任何其他内容。"""



# ========== 字段预览Prompt ==========

PROMPT_FIELD_PREVIEW = """你是医疗数据分析专家。请仔细分析这份医疗文档，识别所有可以提取的数据字段。

特别注意：如果文档是量表、问卷或评分表（如Likert量表、VAS评分等），需要：
- 识别每个评分项目的编号/名称作为独立字段（如"Pre_ICCAS_01"、"PHQ_01"等）
- 观察每个项目的勾选框（checkbox）位置，判断被勾选的是第几个选项
- 根据量表的评分规则（如1-5分、0-3分），将勾选位置映射为对应的数值
- example_value必须是该项目被勾选的实际数值（如勾选第4个框则值为"4"）
- field_type对于评分项应为"number"

返回JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text、number、date",
      "example_value": "从文档中实际提取的真实值",
      "confidence": 0.95,
      "category": "类别"
    }}
  ]
}}

要求:
1. 字段名称使用原文，简洁准确
2. 尽量识别所有数据：基本信息、诊断、检查结果、评分指标、量表项目、药物、治疗方案等
3. category取值范围：基本信息、诊断信息、检查结果、治疗记录、评分量表、问卷项目、其他
4. example_value必须是从文档中实际提取的真实值，对于勾选框请转换为对应数值
5. confidence表示该字段在文档中的识别置信度(0-1)
6. 按category分组，同类别字段放在一起

只返回JSON，不要包含任何其他内容。"""

PROMPT_FIELD_PREVIEW_TEXT = """你是医疗数据分析专家。请仔细分析以下医疗文本，识别所有可以提取的数据字段。

输出JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text或number或date",
      "example_value": "从文本中提取的示例值",
      "confidence": 0.95,
      "category": "类别"
    }}
  ]
}}

要求:
1. 字段名称使用中文，简洁明确
2. 尽可能识别所有有意义的数据项
3. category取值范围：基本信息、检验结果、诊疗记录、护理评估、其他
4. example_value必须是从文本中实际提取的真实值
5. confidence表示该字段的识别置信度(0-1)

文本内容:
{text_content}

只输出JSON，不要输出任何其他内容。"""

PROMPT_EXTRACT_FIELD_NAMES = """你是医疗模板设计专家，当前视角为{role_hint}。请仔细分析以下医疗相关文本，从中识别所有可以作为"数据提取模板字段名"的医学术语或概念。

注意：你要提取的是**字段名称**（如"血压"、"诊断"、"护理措施"），而不是具体的数据值。

输出JSON格式:
{{
  "fields": [
    {{
      "name": "字段名称",
      "category": "类别",
      "confidence": 0.95
    }}
  ]
}}

要求:
1. 字段名称使用中文，简洁明确（2-8个字为宜）
2. 只提取与医疗场景相关的字段，忽略无关词汇
3. category取值范围：基本信息、检验结果、诊疗记录、护理评估、科研数据、其他
4. confidence表示该词作为模板字段的合理程度(0-1)
5. 不要重复提取含义相同的字段
6. 根据{role_hint}的视角，优先识别该角色关注的字段
7. 不要将患者的具体姓名、具体数值等作为字段名

文本内容:
{text_content}

只输出JSON，不要输出任何其他内容。"""




# ========== 模板文件分析Prompt ==========

PROMPT_TEMPLATE_ANALYSIS = """你是医疗表单结构分析专家。你正在查看的是一份**空白模板/问卷/量表**（不是填写好的文档），请识别这份表单设计要收集的所有数据字段。

输出JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text或number或date或score",
      "category": "类别",
      "confidence": 0.95
    }}
  ],
  "template_info": {{
    "suggested_name": "建议的模板名称",
    "is_scoring_scale": false,
    "suggested_department": "general"
  }}
}}

要求:
1. 这是空白表单，关注的是表单**结构**而非填入的数据值
2. 识别所有表头、行标签、列标签、评分项、填写栏位等
3. field_name 使用中文，简洁明确（2-8字）
4. field_type: text(文本)、number(数值)、date(日期)、score(评分项)
5. category取值范围：基本信息、检验结果、诊疗记录、护理评估、量表评分、问卷项目、其他
6. 如果是评分量表（如Barthel、Morse等），设 is_scoring_scale 为 true
7. suggested_department 从以下选择: general, cardiology, neurology, surgery, pediatrics, obstetrics, emergency
8. 按category分组，同类别字段放在一起

只输出JSON，不要输出任何其他内容。"""

PROMPT_TEMPLATE_ANALYSIS_TEXT = """你是医疗表单结构分析专家。以下文本描述了一份医疗文档或量表的结构。请识别其中设计要收集的所有数据字段。

输出JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text或number或date或score",
      "category": "类别",
      "confidence": 0.95
    }}
  ],
  "template_info": {{
    "suggested_name": "建议的模板名称",
    "is_scoring_scale": false,
    "suggested_department": "general"
  }}
}}

要求:
1. 关注文档设计的**结构和字段**，不是具体数据值
2. field_name 使用中文，简洁明确（2-8字）
3. field_type: text(文本)、number(数值)、date(日期)、score(评分项)
4. category取值范围：基本信息、检验结果、诊疗记录、护理评估、量表评分、问卷项目、其他
5. 如果是评分量表，设 is_scoring_scale 为 true
6. suggested_department 从以下选择: general, cardiology, neurology, surgery, pediatrics, obstetrics, emergency

文本内容:
{text_content}

只输出JSON，不要输出任何其他内容。"""
# ========== 音频专用Prompt模板 ==========

PROMPT_AUDIO_DOCTOR = """你是临床医生数据提取专家。以下是医患对话的语音转录文本，请从中提取结构化病历信息。

## 输出格式(JSON):
{
  "patient_info": {"姓名": "", "性别": "", "年龄": null},
  "chief_complaint": "",
  "present_illness": "",
  "past_history": "",
  "physical_exam": {},
  "diagnosis": [{"诊断名称": "", "ICD10编码": ""}],
  "treatment_plan": {"药物治疗": "", "医嘱": "", "其他": ""},
  "conversation_notes": "",
  "confidence": {}
}

## 提取规则:
1. 从对话中识别患者自述的症状和病史
2. 提取医生口述的诊断和治疗建议
3. 主诉通常是患者开场描述的主要不适
4. 注意区分医生询问和患者回答
5. 如信息不完整或无法识别，对应字段留空字符串
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""

PROMPT_AUDIO_NURSE = """你是护理评估专家。以下是护理交班或患者访谈的语音转录文本，请提取护理相关信息。

## 输出格式(JSON):
{
  "patient_info": {"姓名": "", "床号": "", "科室": ""},
  "vital_signs_verbal": {},
  "nursing_observations": "",
  "patient_complaints": "",
  "nursing_actions": "",
  "handover_notes": "",
  "risk_alerts": "",
  "confidence": {}
}

## 提取规则:
1. 识别口述的生命体征数值（体温、血压、脉搏、呼吸、血氧等）
2. 提取护理观察内容（皮肤、伤口、活动能力、意识状态）
3. 记录患者主观感受和主诉
4. 提取交班时的重点提醒事项
5. 识别提及的护理风险（跌倒、压疮、管路等）
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""

PROMPT_AUDIO_RESEARCHER = """你是临床科研数据提取专家。以下是研究访谈或病历口述的语音转录文本，请提取科研相关数据。

## 输出格式(JSON):
{
  "demographics": {"姓名": "", "性别": "", "年龄": null, "职业": "", "教育程度": ""},
  "medical_history": "",
  "intervention_details": "",
  "outcome_measures": "",
  "patient_experience": "",
  "adherence_notes": "",
  "adverse_events": "",
  "research_notes": "",
  "confidence": {}
}

## 提取规则:
1. 提取人口学特征
2. 识别干预措施的描述
3. 提取患者自我报告的结局（症状改善、生活质量变化）
4. 注意提及的依从性和不良反应
5. 日期格式统一为YYYY-MM-DD
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""


# ========== 科室专属Prompt定义 ==========

# --- 心内科 ---
PROMPT_CARDIOLOGY_CLINICAL = """你是心内科数据提取专家。请识别该心内科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "心电图": "",
    "心脏超声": "",
    "EF值": null,
    "BNP_NT_proBNP": null,
    "冠脉造影结果": "",
    "心律失常类型": "",
    "心功能分级_NYHA": "",
    "血压_收缩压_舒张压": "",
    "心率": null,
    "血脂_LDL_HDL_TG_TC": "",
    "肌钙蛋白": "",
    "用药方案": "",
    "PCI_CABG记录": "",
    "诊断": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. EF值提取纯数字(百分比)。BNP/NT-proBNP提取数值及单位。
2. 冠脉造影需详细记录病变血管和狭窄程度。
3. 心功能分级按NYHA I-IV级标准。
4. 用药方案需区分抗血小板、他汀、ACEI/ARB、β受体阻滞剂等类别。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_NURSING = """你是心内科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "心电监护": "",
    "生命体征_体温_脉搏_呼吸_血压": "",
    "胸痛评估": "",
    "出入量记录": "",
    "活动耐量评估": "",
    "抗凝药物管理": "",
    "跌倒风险评估": "",
    "心理状态": "",
    "饮食护理": "",
    "心脏康复指导": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 生命体征需提取完整(体温、脉搏、呼吸、血压)。
2. 胸痛评估包含部位、性质、持续时间、NRS评分。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_OTHER = """你是心血管科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "人口学特征": "",
    "LVEF": null,
    "BNP": null,
    "LDL_C": null,
    "支架类型": "",
    "再狭窄": "",
    "MACE事件": "",
    "随访日期": "",
    "用药情况": "",
    "治疗结局": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_AUDIO = """你是心内科医疗对话分析专家。请从该心内科相关的语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "心电图提及": "",
    "EF值": null,
    "BNP": null,
    "用药方案": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 神经内科 ---
PROMPT_NEUROLOGY_CLINICAL = """你是神经内科数据提取专家。请识别该神经内科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "GCS评分": null,
    "NIHSS评分": null,
    "肌力分级": "",
    "感觉障碍": "",
    "反射检查": "",
    "头颅CT_MRI": "",
    "脑电图": "",
    "腰穿结果": "",
    "病灶定位": "",
    "发病时间": "",
    "溶栓_取栓记录": "",
    "诊断": "",
    "治疗方案": ""
  },
  "confidence": {}
}

## 提取规则
1. GCS评分提取总分(3-15)和各项分值(E/V/M)。NIHSS提取总分(0-42)。
2. 肌力按0-5级标准，需注明部位(左上/左下/右上/右下)。
3. 发病时间精确到小时(用于评估溶栓时间窗)。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_NURSING = """你是神经内科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "意识水平评估": "",
    "生命体征": "",
    "吞咽功能筛查": "",
    "跌倒风险评估": "",
    "压疮风险评估": "",
    "肢体活动度": "",
    "康复活动记录": "",
    "用药护理": "",
    "安全护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 意识水平需包含GCS评分或具体描述。
2. 吞咽筛查结果需明确通过/未通过及分级。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_OTHER = """你是神经科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科"
  },
  "custom_fields": {
    "人口学特征": "",
    "mRS评分": null,
    "NIHSS评分": null,
    "发病至治疗时间": "",
    "梗死体积": "",
    "治疗方式": "",
    "复发情况": "",
    "功能预后": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_AUDIO = """你是神经内科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "GCS评分": null,
    "NIHSS评分": null,
    "影像学提及": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 外科 ---
PROMPT_SURGERY_CLINICAL = """你是外科数据提取专家。请识别该外科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "手术名称": "",
    "ASA分级": "",
    "麻醉方式": "",
    "切口类型": "",
    "术中出血量_ml": null,
    "术中输血": "",
    "手术时长_min": null,
    "引流管情况": "",
    "病理结果": "",
    "术后诊断": "",
    "出院医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 术中出血量和手术时长提取纯数字。
2. ASA分级按I-V级标准。切口类型分清洁/清洁-污染/污染/感染。
3. 引流管需记录类型、位置、引流量。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_NURSING = """你是外科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "术前评估": "",
    "生命体征": "",
    "切口护理": "",
    "引流管护理": "",
    "VTE预防": "",
    "疼痛评分_NRS": null,
    "术后活动": "",
    "饮食护理": "",
    "用药护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 切口护理需包含愈合分级(甲/乙/丙)和敷料情况。
2. 引流管护理需包含引流量、颜色、性状。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_OTHER = """你是外科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科"
  },
  "custom_fields": {
    "人口学特征": "",
    "手术方式": "",
    "手术时长_min": null,
    "出血量_ml": null,
    "并发症": "",
    "住院天数": null,
    "SSI发生": "",
    "再手术": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_AUDIO = """你是外科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "手术方案": "",
    "术后情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 儿科 ---
PROMPT_PEDIATRICS_CLINICAL = """你是儿科数据提取专家。请识别该儿科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "出生体重_g": null,
    "胎龄_周": null,
    "体重_kg": null,
    "身高_cm": null,
    "头围_cm": null,
    "喂养方式": "",
    "疫苗接种史": "",
    "生长发育评估": "",
    "体格检查": "",
    "辅助检查": "",
    "诊断": "",
    "用药_剂量_kg": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. 年龄需区分"岁"和"月龄"(如3岁2月、6月龄)。
2. 体重精确到0.1kg，身高精确到0.1cm。
3. 用药需注明按体重剂量(mg/kg)。
4. 疫苗接种记录需列出已接种和未接种疫苗。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_NURSING = """你是儿科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "体温调节": "",
    "喂养评估": "",
    "FLACC疼痛评分": null,
    "皮肤评估": "",
    "生命体征": "",
    "安全评估": "",
    "家长宣教": "",
    "排泄记录": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. FLACC疼痛评分适用于无法自述疼痛的儿童(0-10分)。
2. 喂养评估需包含奶量/辅食量、喂养耐受情况。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_OTHER = """你是儿科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科"
  },
  "custom_fields": {
    "人口学特征": "",
    "身高百分位": "",
    "体重百分位": "",
    "BMI_Z评分": null,
    "发育里程碑": "",
    "营养评估": "",
    "疫苗完成率": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 百分位数需标注参考标准(WHO/CDC)。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_AUDIO = """你是儿科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "喂养情况": "",
    "发育情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 妇产科 ---
PROMPT_OBSTETRICS_CLINICAL = """你是妇产科数据提取专家。请识别该妇产科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "孕周": "",
    "孕次产次_GPAL": "",
    "末次月经": "",
    "胎心率": null,
    "宫高_cm": null,
    "腹围_cm": null,
    "B超_胎儿": "",
    "羊水指数": null,
    "血压": "",
    "尿蛋白": "",
    "分娩方式": "",
    "会阴状况": "",
    "Apgar评分": "",
    "新生儿体重_g": null,
    "诊断": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. 孕周格式如"38+2周"。GPAL格式如"G2P1A0L1"。
2. Apgar评分需记录1分钟和5分钟值(如9-10-10)。
3. 胎心率正常范围110-160次/分。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_NURSING = """你是妇产科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "产程监测": "",
    "宫缩评估": "",
    "胎心监护": "",
    "生命体征": "",
    "恶露观察": "",
    "母乳喂养评估": "",
    "产后出血评估": "",
    "会阴护理": "",
    "心理护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 恶露需记录颜色、量、气味。
2. 产后出血评估需包含出血量估计。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_OTHER = """你是妇产科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科"
  },
  "custom_fields": {
    "人口学特征": "",
    "分娩孕周": "",
    "出生体重_g": null,
    "分娩方式": "",
    "产后出血量_ml": null,
    "并发症": "",
    "NICU入住": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_AUDIO = """你是妇产科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科"
  },
  "custom_fields": {
    "主诉": "",
    "孕周": "",
    "症状描述": "",
    "胎心情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 急诊科 ---
PROMPT_EMERGENCY_CLINICAL = """你是急诊科数据提取专家。请识别该急诊医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "分诊级别": "",
    "到达方式": "",
    "发病时间": "",
    "发病至就诊时间": "",
    "主诉": "",
    "MEWS_NEWS评分": null,
    "生命体征": "",
    "急救措施": "",
    "用药记录": "",
    "检查结果": "",
    "会诊记录": "",
    "绿色通道": "",
    "诊断": "",
    "转归": ""
  },
  "confidence": {}
}

## 提取规则
1. 分诊级别按I-IV级标准(I级濒危/II级危重/III级急症/IV级非急症)。
2. 发病至就诊时间精确到小时或分钟。
3. 转归需明确:留观/住院/出院/转院/死亡。
4. MEWS评分提取总分(0-14)。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_NURSING = """你是急诊护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "ABCDE快速评估": "",
    "生命体征动态": "",
    "急救用药记录": "",
    "管路记录": "",
    "标本送检": "",
    "交接班核查": "",
    "疼痛评估": "",
    "安全护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. ABCDE评估包含气道/呼吸/循环/意识/暴露五项。
2. 管路记录需包含管路类型、置管时间、通畅情况。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_OTHER = """你是急诊科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科"
  },
  "custom_fields": {
    "人口学特征": "",
    "就诊至处置时间_min": null,
    "分诊级别": "",
    "分诊准确率": "",
    "留观时长_h": null,
    "非计划重返": "",
    "转归": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_AUDIO = """你是急诊科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科"
  },
  "custom_fields": {
    "主诉": "",
    "发病时间": "",
    "症状描述": "",
    "急救措施": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 科室自动检测Prompt ---
PROMPT_DETECT_DEPARTMENT = """你是医疗科室分类专家。请分析以下医疗文本，判断它最可能属于哪个临床科室，并推荐合适的文档类型。

## 可选科室
1. cardiology (心内科) - 特征: 心电图、冠脉造影、EF值、BNP、NT-proBNP、心功能、PCI、房颤、心梗、心律失常、支架
2. neurology (神经内科) - 特征: GCS、NIHSS、偏瘫、脑梗、脑出血、癫痫、脑电图、腰穿、溶栓、卒中
3. surgery (外科) - 特征: 手术、切口、引流、麻醉、ASA、病理、术后、缝合、腹腔镜、清创
4. pediatrics (儿科) - 特征: 月龄、新生儿、疫苗、生长发育、喂养、辅食、百分位、胎龄
5. obstetrics (妇产科) - 特征: 孕周、产检、胎心、宫缩、分娩、剖宫产、恶露、Apgar、GPAL
6. emergency (急诊科) - 特征: 分诊、抢救、急救、120、绿色通道、MEWS、CPR、电除颤
7. general (通用) - 以上均不明确匹配时选择

## 文档类型
- clinical: 诊疗文档(病历、检查报告、诊断、治疗记录)
- nursing: 护理文档(护理评估、量表、护理记录)
- other: 其他(科研数据、访谈、综合报告)

## 输出格式
请直接输出合法的JSON对象：
{
  "department": "科室英文ID",
  "confidence": 0.0,
  "reasoning": "判断依据(30字以内)",
  "sub_type": "clinical或nursing或other"
}

只输出JSON，不要输出任何其他内容。

## 待分析文本：
"""

# 科室Prompt映射表（用于模板初始化和查找）
DEPARTMENT_PROMPTS = {
    'cardiology': {
        'clinical': PROMPT_CARDIOLOGY_CLINICAL,
        'nursing': PROMPT_CARDIOLOGY_NURSING,
        'other': PROMPT_CARDIOLOGY_OTHER,
        'audio': PROMPT_CARDIOLOGY_AUDIO,
    },
    'neurology': {
        'clinical': PROMPT_NEUROLOGY_CLINICAL,
        'nursing': PROMPT_NEUROLOGY_NURSING,
        'other': PROMPT_NEUROLOGY_OTHER,
        'audio': PROMPT_NEUROLOGY_AUDIO,
    },
    'surgery': {
        'clinical': PROMPT_SURGERY_CLINICAL,
        'nursing': PROMPT_SURGERY_NURSING,
        'other': PROMPT_SURGERY_OTHER,
        'audio': PROMPT_SURGERY_AUDIO,
    },
    'pediatrics': {
        'clinical': PROMPT_PEDIATRICS_CLINICAL,
        'nursing': PROMPT_PEDIATRICS_NURSING,
        'other': PROMPT_PEDIATRICS_OTHER,
        'audio': PROMPT_PEDIATRICS_AUDIO,
    },
    'obstetrics': {
        'clinical': PROMPT_OBSTETRICS_CLINICAL,
        'nursing': PROMPT_OBSTETRICS_NURSING,
        'other': PROMPT_OBSTETRICS_OTHER,
        'audio': PROMPT_OBSTETRICS_AUDIO,
    },
    'emergency': {
        'clinical': PROMPT_EMERGENCY_CLINICAL,
        'nursing': PROMPT_EMERGENCY_NURSING,
        'other': PROMPT_EMERGENCY_OTHER,
        'audio': PROMPT_EMERGENCY_AUDIO,
    },
}

# 科室子模板名称映射
DEPT_SUB_NAMES = {
    'clinical': '诊疗记录',
    'nursing': '护理评估',
    'other': '科研数据',
    'audio': '语音录入',
}

# ========== ICCAS 问卷专用 Prompt ==========

PROMPT_ICCAS_QUESTIONNAIRE = """You are a specialist in reading scanned paper questionnaires. This image is a printed ICCAS (Interprofessional Collaborative Competency Attainment Survey) questionnaire that has been filled in by hand with checkmarks/ticks in checkbox squares.

## Your Task
Visually inspect every checkbox area in the image. A box with ANY mark inside (tick, cross, filled) counts as CHECKED. An empty box counts as UNCHECKED. Output results in the EXACT order they appear on the questionnaire. Do NOT translate, rewrite, or abbreviate any text — output the original English wording from the questionnaire.

## Recognition Rules
- Single-choice only: if a row has multiple boxes, only ONE should be checked. Pick the one with the clearest mark.
- If NO box is checked for a question, output "Not answered".
- For "Year of study": recognise the handwritten digit OR the checked digit box.
- PART 1 (PRE) and PART 2 (POST) are completely separate sections. Never mix them.

## Required JSON Output Structure
Return ONLY the following JSON object, no markdown, no explanation:

{
  "Consent": {
    "I am now 18 years old or above": 0,
    "I identify as a student enrolled in the course Advising on Smoking Cessation": 0,
    "I have read and understood the objectives and procedures of this research project": 0,
    "I consent to participate in this survey and understand that I can withdraw at any time during the research period": 0
  },
  "PRE_SURVEY": {
    "Demographics": {
      "Discipline": "",
      "Year_of_study": null,
      "Gender": "",
      "Prior_smoking_cessation_training": "",
      "Smoking_status": ""
    },
    "ICCAS_PRE": {
      "Pre_ICCAS_01": null,
      "Pre_ICCAS_02": null,
      "Pre_ICCAS_03": null,
      "Pre_ICCAS_04": null,
      "Pre_ICCAS_05": null,
      "Pre_ICCAS_06": null,
      "Pre_ICCAS_07": null,
      "Pre_ICCAS_08": null,
      "Pre_ICCAS_09": null,
      "Pre_ICCAS_10": null,
      "Pre_ICCAS_11": null,
      "Pre_ICCAS_12": null,
      "Pre_ICCAS_13": null,
      "Pre_ICCAS_14": null,
      "Pre_ICCAS_15": null,
      "Pre_ICCAS_16": null,
      "Pre_ICCAS_17": null,
      "Pre_ICCAS_18": null,
      "Pre_ICCAS_19": null,
      "Pre_ICCAS_20": null
    },
    "KAP_PRE": {
      "Knowledge_01": "",
      "Knowledge_02": "",
      "Knowledge_03": "",
      "Knowledge_04": "",
      "Knowledge_05": "",
      "Attitudes_01": "",
      "Attitudes_02": "",
      "Attitudes_03": "",
      "Attitudes_04": "",
      "Attitudes_05": "",
      "Practices_01": null,
      "Practices_02": null,
      "Practices_03": null,
      "Practices_04": null,
      "Practices_05": null,
      "Practices_06": null,
      "Practices_07": null,
      "Practices_08": null,
      "Practices_09": null,
      "Practices_10": null,
      "Practices_11": null,
      "Practices_12": null,
      "Practices_13": null
    }
  },
  "POST_SURVEY": {
    "ICCAS_POST": {
      "Post_ICCAS_01": null,
      "Post_ICCAS_02": null,
      "Post_ICCAS_03": null,
      "Post_ICCAS_04": null,
      "Post_ICCAS_05": null,
      "Post_ICCAS_06": null,
      "Post_ICCAS_07": null,
      "Post_ICCAS_08": null,
      "Post_ICCAS_09": null,
      "Post_ICCAS_10": null,
      "Post_ICCAS_11": null,
      "Post_ICCAS_12": null,
      "Post_ICCAS_13": null,
      "Post_ICCAS_14": null,
      "Post_ICCAS_15": null,
      "Post_ICCAS_16": null,
      "Post_ICCAS_17": null,
      "Post_ICCAS_18": null,
      "Post_ICCAS_19": null,
      "Post_ICCAS_20": null,
      "Post_ICCAS_21": ""
    },
    "KAP_POST": {
      "Knowledge_01": "",
      "Knowledge_02": "",
      "Knowledge_03": "",
      "Knowledge_04": "",
      "Knowledge_05": "",
      "Attitudes_01": "",
      "Attitudes_02": "",
      "Attitudes_03": "",
      "Attitudes_04": "",
      "Attitudes_05": "",
      "Practices_01": null,
      "Practices_02": null,
      "Practices_03": null,
      "Practices_04": null,
      "Practices_05": null,
      "Practices_06": null,
      "Practices_07": null,
      "Practices_08": null,
      "Practices_09": null,
      "Practices_10": null,
      "Practices_11": null,
      "Practices_12": null,
      "Practices_13": null
    }
  }
}

## Field Value Rules (STRICT)

### Consent (4 items)
Value: 1 if checked, 0 if unchecked.

### Demographics
- Discipline: output the checked option text exactly as printed (MBBS / Bpharm / BchinMed / Bnurs / BDS). "Not answered" if none checked.
- Year_of_study: integer (handwritten digit or checked digit). null if unreadable.
- Gender: "Female" / "Male" / "Prefer not to say". "Not answered" if none checked.
- Prior_smoking_cessation_training: "Yes" / "No". "Not answered" if none checked.
- Smoking_status: "Non-smoker" / "Ex-smoker (>6 months)" / "Smoker/current (<6 months)". "Not answered" if none checked.

### ICCAS PRE (items 1–20) and ICCAS POST (items 1–20)
Integer 1–5 based on which column checkbox is checked (leftmost = 1, rightmost = 5). null if not answered.
The 5-point scale maps checkbox positions left-to-right: 1 = Strongly Disagree, 2 = Disagree, 3 = Neutral, 4 = Agree, 5 = Strongly Agree.

### ICCAS POST item 21
Output the checked option text exactly: "Much better now" / "Somewhat better now" / "About the same" / "Somewhat worse now" / "Much worse now". "Not answered" if none checked.

### KAP Knowledge (items 1–5, both PRE and POST)
Output: "True" or "False". "Not answered" if none checked.

### KAP Attitudes (items 1–5, both PRE and POST)
Output the checked option text exactly: "Strongly Agree" / "Agree" / "Neutral" / "Disagree" / "Strongly Disagree". "Not answered" if none checked.

### KAP Practices/Confidence (items 1–13, both PRE and POST)
Integer 1–5 based on which column checkbox is checked (leftmost = 1, rightmost = 5). null if not answered.

## Critical Reminders
1. Output ONLY the JSON object above. No markdown fences, no commentary.
2. Every field must be present in the output even if "Not answered" or null.
3. PRE and POST sections are independent — never copy values between them.
4. Maintain exact questionnaire order — do not rearrange items.
5. Carefully examine each checkbox area for marks — even faint ticks count as checked."""


# 科室子模板display_layout映射
DEPT_SUB_LAYOUTS = {
    'clinical': 'table',
    'nursing': 'card',
    'other': 'table',
    'audio': 'table',
}




# ========== 质性研究分析（科研角色专用） ==========

PROMPT_QUALITATIVE_ANALYSIS = """你是临床定性研究专家。请对以下医疗访谈转录文本进行定性分析。

输出JSON格式:
{{
  "themes": ["主题1", "主题2"],
  "keywords": ["关键词1", "关键词2"],
  "codes": [{{"code": "编码类别", "segments": ["相关文本片段1", "片段2"]}}],
  "sentiment": "积极/中性/消极",
  "summary": "2-3句话分析总结"
}}

分析要求:
1. 主题分析: 识别3-5个核心讨论主题
2. 关键词提取: 提取10-15个关键医学/情感词汇
3. 编码分类: 按类别(如症状描述、治疗态度、医患沟通、情感表达、生活影响)编码文本
4. 情感倾向: 判断整体情感

转录文本:
{transcript}

只输出JSON，不要输出任何其他内容。"""


def qualitative_analysis(transcript_text):
    """对转录文本进行定性研究分析（发送前自动脱敏）"""
    masked_text, _report = desensitize_text(transcript_text)
    prompt = PROMPT_QUALITATIVE_ANALYSIS.format(transcript=masked_text)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=3000
    )
    return parse_ai_response(response.choices[0].message.content)


# ========== 文本文件解析与预处理 ==========

def _parse_text_file(file_path):
    """解析txt/docx文件为纯文本"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise Exception("无法识别文本文件编码，请使用UTF-8编码保存")

    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise Exception("python-docx 库未安装，请运行 pip install python-docx")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)

    elif ext == '.doc':
        raise Exception("不支持旧版.doc格式，请将文件另存为.docx后重新上传")

    else:
        raise Exception(f"不支持的文本格式: {ext}")


def _preprocess_text(text):
    """文本预处理：标准化格式"""
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ========== 增强版质性研究分析 ==========

QUALITATIVE_TYPE_HINTS = {
    'interview': {
        'cn': '深度访谈',
        'hints': '患者体验、疾病认知、治疗态度、就医过程、心理感受',
        'coding': '症状描述、情感表达、医患沟通、治疗依从性、生活影响'
    },
    'focus_group': {
        'cn': '焦点小组',
        'hints': '群体共识、争议点、互动模式、观点演变、关键事件',
        'coding': '观点类别、互动类型、共识差异、关键事件、群体动态'
    },
    'observation': {
        'cn': '观察记录',
        'hints': '行为模式、环境因素、非语言信息、事件序列、场景特征',
        'coding': '行为类型、场景因素、时间特征、主体角色、环境条件'
    }
}

PROMPT_QUALITATIVE_ENHANCED = """你是质性研究方法论专家。请对以下{analysis_type_cn}材料严格按照四步质性分析法进行系统分析。

## 分析步骤

### 第一步：初始编码（Open Coding）
逐句或逐段阅读文本，识别有意义的概念、想法、行为模式，并为每个有意义的片段打上编码标签。
- 每个编码包含唯一编号(C01, C02...)、编码标签（简短概念名）、对应的原文片段、段落编号(P1, P2...)

### 第二步：主题聚类（Theme Clustering）
将相似的初始编码进行归纳合并，形成更高层次的上位主题类别。
- 识别3-5个主主题
- 每个主主题下有1-4个子主题
- 每个子主题关联具体的code_id列表
- 确保主题内部逻辑一致性和主题间差异性

### 第三步：典型原话保留（Representative Quotes）
为每个最终确定的主题选择2-3条最具代表性的原始引语。
- 必须是原文直接引用，不做改编
- 应能充分支撑该主题的核心观点

### 第四步：层级化输出（Hierarchical Structure）
最终输出严格按照"主题—子主题—编码—原话摘录"的层级结构。

## 输出JSON格式（严格遵守此结构）
{{
  "methodology_note": "本分析采用{analysis_type_cn}质性研究方法，遵循开放性编码→主题聚类→代表性引用→层级输出的四步分析流程",
  "analysis_type": "{analysis_type}",
  "step1_initial_coding": [
    {{
      "code_id": "C01",
      "code_label": "编码标签名",
      "original_text": "原文中的具体片段",
      "paragraph_ref": "P1"
    }}
  ],
  "step2_theme_clustering": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "sub_theme": "子主题名称",
          "codes": ["C01", "C03"],
          "description": "该子主题的简要描述"
        }}
      ]
    }}
  ],
  "step3_representative_quotes": [
    {{
      "theme": "主主题名称",
      "quotes": [
        "原话引用1",
        "原话引用2"
      ]
    }}
  ],
  "step4_hierarchical_output": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "name": "子主题名称",
          "codes": [
            {{
              "label": "编码标签",
              "quotes": ["支撑该编码的原话"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

## 分析要点
- 分析关注点：{analysis_hints}
- 编码参考类别：{coding_categories}
- 初始编码数量：8-20个（视文本长度而定）
- 主主题数量：3-5个
- 每个主题下的代表性引用：2-3条
- 所有引用必须来自原文，保持原始措辞

## 待分析文本：
{transcript}

只输出JSON，不要输出任何其他内容。"""


