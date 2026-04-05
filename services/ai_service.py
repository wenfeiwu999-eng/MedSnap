# -*- coding: utf-8 -*-
"""AI recognition, audio transcription, qualitative analysis"""

import os, json, re, base64
from PIL import Image
from config import (
    client, MODEL_NAME, DASHSCOPE_API_KEY,
    HAS_TESSERACT, HAS_PYMUPDF, HAS_DASHSCOPE, HAS_PYDUB,
)
from desensitizer import desensitize_text, desensitize_structured_data
from services.file_processor import preprocess_image, image_to_base64
from prompts import PROMPT_QUALITATIVE_ANALYSIS, PROMPT_QUALITATIVE_ENHANCED, QUALITATIVE_TYPE_HINTS

# ========== AI 识别核心 ==========

def extract_from_ocr_text(ocr_text, ai_prompt):
    """将本地OCR提取的文本发送给LLM进行结构化提取（发送前自动脱敏）"""
    # 脱敏：在发送远程LLM之前，对OCR文本进行敏感信息脱敏
    masked_text, _report = desensitize_text(ocr_text)
    combined_prompt = ai_prompt + "\n\n以下是通过OCR识别出的医疗文档文本，请按上述要求提取结构化信息：\n\n" + masked_text
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': combined_prompt}],
        temperature=0.1,
        max_tokens=16384
    )
    raw_text = response.choices[0].message.content
    parsed = parse_ai_response(raw_text)
    return parsed, raw_text


def extract_medical_data_multimodal(image_path, ai_prompt):
    """调用多模态模型直接识别图片（原始方式，作为OCR失败时的回退）
    注意：多模态模式下图片会发送至远程AI服务，返回的结构化数据将在本地进行脱敏。"""
    print("[PRIVACY] 多模态模式：图片将发送至AI服务进行识别")
    preprocessed_path = preprocess_image(image_path)
    b64_image = image_to_base64(preprocessed_path)

    ext = os.path.splitext(image_path)[1].lower()
    mime_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.bmp': 'image/bmp',
        '.tiff': 'image/tiff'
    }
    mime_type = mime_map.get(ext, 'image/jpeg')

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{
            'role': 'user',
            'content': [
                {'type': 'text', 'text': ai_prompt},
                {
                    'type': 'image_url',
                    'image_url': {
                        'url': f"data:{mime_type};base64,{b64_image}"
                    }
                }
            ]
        }],
        temperature=0.1,
        max_tokens=16384
    )

    raw_text = response.choices[0].message.content

    try:
        if preprocessed_path != image_path and os.path.exists(preprocessed_path):
            os.remove(preprocessed_path)
    except Exception:
        pass

    parsed = parse_ai_response(raw_text)
    # 脱敏：对多模态返回的结构化数据进行脱敏
    parsed = desensitize_structured_data(parsed)
    return parsed, raw_text


def extract_medical_data(image_path, ai_prompt):
    """图片识别调度器：优先本地OCR + LLM结构化，失败则回退到多模态直接识别"""
    # 策略1: 尝试本地OCR
    if HAS_TESSERACT:
        try:
            ocr_text = local_ocr(image_path)
            if ocr_text and len(ocr_text) >= 10:
                print(f"[OCR] 本地OCR成功，提取文本长度: {len(ocr_text)}")
                parsed, raw_text = extract_from_ocr_text(ocr_text, ai_prompt)
                if 'error' not in parsed:
                    return parsed, raw_text
                print(f"[OCR] OCR文本结构化失败，回退到多模态识别")
            else:
                print(f"[OCR] 本地OCR文本过短({len(ocr_text) if ocr_text else 0}字符)，回退到多模态识别")
        except Exception as e:
            print(f"[OCR] 本地OCR失败: {e}，回退到多模态识别")

    # 策略2: 多模态模型直接识别
    return extract_medical_data_multimodal(image_path, ai_prompt)


def parse_ai_response(raw_text):
    text = re.sub(r'<think>.*?</think>', '', raw_text, flags=re.DOTALL)
    text = text.strip()
    if '```' in text:
        match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
        if match:
            text = match.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass
    return {"error": "AI返回结果解析失败，请重试或检查图片质量", "raw_response": raw_text[:500]}


# ========== 语音识别与文本分析 ==========

def transcribe_audio(audio_path):
    """调用DashScope Paraformer识别本地音频文件"""
    if not HAS_DASHSCOPE:
        raise Exception("dashscope 库未安装，请运行 pip install dashscope")
    if not DASHSCOPE_API_KEY:
        raise Exception("环境变量 DASHSCOPE_API_KEY 未设置，无法使用语音识别")

    ext = os.path.splitext(audio_path)[1].lower()
    actual_path = audio_path
    converted = False

    # 不支持的格式转为wav
    if ext in ('.m4a', '.flac'):
        if not HAS_PYDUB:
            raise Exception("pydub 库未安装，无法转换 m4a/flac 格式，请运行 pip install pydub")
        audio_seg = AudioSegment.from_file(audio_path)
        actual_path = audio_path + '.wav'
        audio_seg.export(actual_path, format='wav')
        converted = True

    fmt_map = {'.wav': 'wav', '.mp3': 'mp3', '.aac': 'aac',
               '.amr': 'amr', '.opus': 'opus'}
    fmt = fmt_map.get(os.path.splitext(actual_path)[1].lower(), 'wav')

    try:
        recognition = Recognition(
            model='paraformer-realtime-v2',
            format=fmt,
            sample_rate=16000,
            language_hints=['zh', 'en']
        )
        result = recognition.call(actual_path)

        # 提取完整文本
        sentences = []
        if hasattr(result, 'get_sentence') and callable(result.get_sentence):
            sentences = result.get_sentence() or []
        full_text = ''.join([s.get('text', '') for s in sentences]) if sentences else ''

        if not full_text:
            # 尝试从output中获取
            if hasattr(result, 'output') and result.output:
                out = result.output
                if isinstance(out, dict) and 'text' in out:
                    full_text = out['text']
                elif isinstance(out, dict) and 'sentence' in out:
                    for s in out['sentence']:
                        full_text += s.get('text', '')

        if not full_text:
            raise Exception("语音识别未返回有效文本，请检查音频文件质量")

        return {
            'text': full_text,
            'sentences': sentences,
            'language': 'zh'
        }
    finally:
        if converted and os.path.exists(actual_path):
            try:
                os.remove(actual_path)
            except Exception:
                pass


def extract_from_transcript(transcript_text, ai_prompt):
    """用Qwen模型从转录文本中提取结构化数据（纯文本模式，发送前自动脱敏）"""
    # 脱敏：在发送远程LLM之前，对转录文本进行敏感信息脱敏
    masked_text, _report = desensitize_text(transcript_text)
    combined_prompt = ai_prompt + "\n\n以下是语音转录文本，请按上述要求提取结构化信息：\n\n" + masked_text
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': combined_prompt}],
        temperature=0.1,
        max_tokens=16384
    )
    raw_text = response.choices[0].message.content
    parsed = parse_ai_response(raw_text)
    return parsed, raw_text


# ========== 质性研究分析（科研角色专用） ==========

PROMPT_QUALITATIVE_ANALYSIS = """你是临床定性研究专家。请对以下医疗访谈转录文本进行定性分析。

输出JSON格式:
{{
  "themes": ["主题1", "主题2"],
  "keywords": ["关键词1", "关键词2"],
  "codes": [{{"code": "编码类别", "segments": ["相关文本片段1", "片段2"]}}],
  "sentiment": "积极/中性/消极",
  "summary": "2-3句话分析总结"
}}

分析要求:
1. 主题分析: 识别3-5个核心讨论主题
2. 关键词提取: 提取10-15个关键医学/情感词汇
3. 编码分类: 按类别(如症状描述、治疗态度、医患沟通、情感表达、生活影响)编码文本
4. 情感倾向: 判断整体情感

转录文本:
{transcript}

只输出JSON，不要输出任何其他内容。"""


def qualitative_analysis(transcript_text):
    """对转录文本进行定性研究分析（发送前自动脱敏）"""
    masked_text, _report = desensitize_text(transcript_text)
    prompt = PROMPT_QUALITATIVE_ANALYSIS.format(transcript=masked_text)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=3000
    )
    return parse_ai_response(response.choices[0].message.content)


# ========== 文本文件解析与预处理 ==========

def _parse_text_file(file_path):
    """解析txt/docx文件为纯文本"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise Exception("无法识别文本文件编码，请使用UTF-8编码保存")

    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise Exception("python-docx 库未安装，请运行 pip install python-docx")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)

    elif ext == '.doc':
        raise Exception("不支持旧版.doc格式，请将文件另存为.docx后重新上传")

    else:
        raise Exception(f"不支持的文本格式: {ext}")


def _preprocess_text(text):
    """文本预处理：标准化格式"""
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ========== 增强版质性研究分析 ==========

QUALITATIVE_TYPE_HINTS = {
    'interview': {
        'cn': '深度访谈',
        'hints': '患者体验、疾病认知、治疗态度、就医过程、心理感受',
        'coding': '症状描述、情感表达、医患沟通、治疗依从性、生活影响'
    },
    'focus_group': {
        'cn': '焦点小组',
        'hints': '群体共识、争议点、互动模式、观点演变、关键事件',
        'coding': '观点类别、互动类型、共识差异、关键事件、群体动态'
    },
    'observation': {
        'cn': '观察记录',
        'hints': '行为模式、环境因素、非语言信息、事件序列、场景特征',
        'coding': '行为类型、场景因素、时间特征、主体角色、环境条件'
    }
}

PROMPT_QUALITATIVE_ENHANCED = """你是质性研究方法论专家。请对以下{analysis_type_cn}材料严格按照四步质性分析法进行系统分析。

## 分析步骤

### 第一步：初始编码（Open Coding）
逐句或逐段阅读文本，识别有意义的概念、想法、行为模式，并为每个有意义的片段打上编码标签。
- 每个编码包含唯一编号(C01, C02...)、编码标签（简短概念名）、对应的原文片段、段落编号(P1, P2...)

### 第二步：主题聚类（Theme Clustering）
将相似的初始编码进行归纳合并，形成更高层次的上位主题类别。
- 识别3-5个主主题
- 每个主主题下有1-4个子主题
- 每个子主题关联具体的code_id列表
- 确保主题内部逻辑一致性和主题间差异性

### 第三步：典型原话保留（Representative Quotes）
为每个最终确定的主题选择2-3条最具代表性的原始引语。
- 必须是原文直接引用，不做改编
- 应能充分支撑该主题的核心观点

### 第四步：层级化输出（Hierarchical Structure）
最终输出严格按照"主题—子主题—编码—原话摘录"的层级结构。

## 输出JSON格式（严格遵守此结构）
{{
  "methodology_note": "本分析采用{analysis_type_cn}质性研究方法，遵循开放性编码→主题聚类→代表性引用→层级输出的四步分析流程",
  "analysis_type": "{analysis_type}",
  "step1_initial_coding": [
    {{
      "code_id": "C01",
      "code_label": "编码标签名",
      "original_text": "原文中的具体片段",
      "paragraph_ref": "P1"
    }}
  ],
  "step2_theme_clustering": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "sub_theme": "子主题名称",
          "codes": ["C01", "C03"],
          "description": "该子主题的简要描述"
        }}
      ]
    }}
  ],
  "step3_representative_quotes": [
    {{
      "theme": "主主题名称",
      "quotes": [
        "原话引用1",
        "原话引用2"
      ]
    }}
  ],
  "step4_hierarchical_output": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "name": "子主题名称",
          "codes": [
            {{
              "label": "编码标签",
              "quotes": ["支撑该编码的原话"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

## 分析要点
- 分析关注点：{analysis_hints}
- 编码参考类别：{coding_categories}
- 初始编码数量：8-20个（视文本长度而定）
- 主主题数量：3-5个
- 每个主题下的代表性引用：2-3条
- 所有引用必须来自原文，保持原始措辞

## 待分析文本：
{transcript}

只输出JSON，不要输出任何其他内容。"""


def qualitative_analysis_enhanced(transcript_text, analysis_type='interview'):
    """四步法质性分析：初始编码→主题聚类→代表性引用→层级输出（发送前自动脱敏）"""
    masked_text, _report = desensitize_text(transcript_text)
    type_info = QUALITATIVE_TYPE_HINTS.get(analysis_type, QUALITATIVE_TYPE_HINTS['interview'])
    prompt = PROMPT_QUALITATIVE_ENHANCED.format(
        analysis_type_cn=type_info['cn'],
        analysis_type=analysis_type,
        analysis_hints=type_info['hints'],
        coding_categories=type_info['coding'],
        transcript=masked_text
    )
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=4000
    )
    result = parse_ai_response(response.choices[0].message.content)
    # 基本校验：确保返回4步结构
    if 'error' not in result and 'step1_initial_coding' not in result:
        # 可能是旧格式返回，尝试包装
        if 'themes' in result:
            return result  # 返回旧格式，前端会兼容处理
    return result


# ========== 数据分析模块 ==========


